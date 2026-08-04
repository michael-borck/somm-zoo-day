#!/usr/bin/env python3
"""Package the material for Luke into per-event bundles.

Seven loose attachments is a wall. Grouping them by event, with a short cover
note inside each saying what the files are and what is actually being asked,
means the reasoning sits next to the file it is about rather than four
paragraphs earlier in an email.

The coverage map stays loose and unzipped - it is the one thing he is being
asked to open and mark up, so it should not be behind a click.

Writes to bundles/ (gitignored: these carry Luke's slides).
"""

import pathlib
import shutil
import zipfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parent.parent
OUT = ROOT / "bundles"

TEAL = RGBColor(0x1B, 0x3A, 0x4B)

BUNDLES = [
    {
        "dir": "1_Thursday_Showcase",
        "title": "Thursday 20 August — Assessment 2030 Showcase",
        "subtitle": "30 minutes, the two of us. Read this first.",
        "files": [
            "UDL_Lens_Showcase_SUGGESTION_v1.pptx",
            "UDL_Lens_Showcase_AI_half_v1.pptx",
        ],
        "sections": [
            ("What is in here", [
                "UDL_Lens_Showcase_SUGGESTION_v1.pptx — a suggested 15-slide cut of the whole "
                "talk. Five of these slides are yours and completely unchanged; I deleted "
                "rather than rewrote, so nothing has been reworded behind your back.",
                "UDL_Lens_Showcase_AI_half_v1.pptx — my nine slides on their own, in case you "
                "would rather assemble the running order differently.",
            ]),
            ("Why fifteen", [
                "It is arithmetic rather than an opinion about your slides. Thirty minutes, "
                "less five for questions and five for the live demonstration, leaves about "
                "twenty minutes of slides. Your 39 plus my 10 is roughly 24 seconds each.",
                "Nothing is deleted. The ACF and iSoLT depth moves to the handout in the Friday "
                "bundle, where it does more work for someone writing an application than it "
                "ever did on a screen.",
            ]),
            ("One slide you have not seen before", [
                "Slide 8 is a live audience poll — four UDL considerations, one of which is "
                "invented. It runs in about two minutes and sets up the wrong-codes story "
                "rather than just illustrating it. Happy to drop it if you would rather not.",
            ]),
            ("What I would like from you", [
                "Does the fifteen-slide shape work, and is the cut of your slides fair? If you "
                "would rather keep more of yours and cut more of mine, that is genuinely fine "
                "— the target is about fifteen slides, not a particular fifteen.",
            ]),
        ],
    },
    {
        "dir": "2_Friday_Retreat",
        "title": "Friday 21 August — SoMM retreat",
        "subtitle": "Universal Design breakout, Session One, 9.45–10.30am, max 10 people.",
        "files": [
            "UDL_Lens_ZooDay_v1.pptx",
            "UDL_Lens_ZooDay_Facilitator_Pack.docx",
            "UDL_Lens_iSoLT_Handout_v2.docx",
            "UDL_Lens_ZooDay_Print_Materials.docx",
        ],
        "sections": [
            ("What is in here", [
                "UDL_Lens_ZooDay_v1.pptx — seven slides plus appendices. The session is mostly "
                "them working on their own unit, so the deck stays out of the way.",
                "UDL_Lens_ZooDay_Facilitator_Pack.docx — the run sheet, minute by minute, for "
                "the two of us. The room never sees it.",
                "UDL_Lens_iSoLT_Handout_v2.docx — three pages carrying the ACF and iSoLT "
                "detail. Goes on the table at 0:33, not earlier.",
                "UDL_Lens_ZooDay_Print_Materials.docx — rating cards, three sample briefs and "
                "the capture sheet. I will print these; nothing needed from you.",
            ]),
            ("Your two segments", [
                "0:33, five minutes — That PDF is evidence. Four ACF benchmarks.",
                "0:38, four minutes — The cycle. Audit, act, re-audit, document.",
                "Both are in the run sheet with notes on what to do rather than what to say. "
                "Worth a look before we lock it, in case I have you doing something you would "
                "rather not.",
            ]),
            ("The one thing I need checked", [
                "The iSoLT points table in the handout is transcribed straight from your "
                "slides. I have not verified it against the current Activity Planner, so "
                "please read those numbers before it goes in front of anyone.",
            ]),
            ("How the session runs", [
                "The room rates a real assessment brief on paper first, in threes, before the "
                "tool says anything. Then the groups read their ratings out loud, and only "
                "then do we show what the tool said. The gap between the two is the content — "
                "which is why the order matters and why the tool goes second.",
            ]),
        ],
    },
]


def style_document(doc: Document) -> None:
    """House style, and the bits that make it readable by a screen reader.

    Real heading styles rather than bold body text, an explicit document
    language so a screen reader picks the right voice, and 11pt minimum.
    """
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    lang = rpr.makeelement(qn("w:lang"), {qn("w:val"): "en-AU"})
    rpr.append(lang)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15


def write_cover(bundle: dict, path: pathlib.Path) -> None:
    doc = Document()
    style_document(doc)

    h = doc.add_heading(bundle["title"], level=1)
    for run in h.runs:
        run.font.color.rgb = TEAL

    sub = doc.add_paragraph(bundle["subtitle"])
    sub.runs[0].italic = True

    for heading, paragraphs in bundle["sections"]:
        h2 = doc.add_heading(heading, level=2)
        for run in h2.runs:
            run.font.color.rgb = TEAL
        for text in paragraphs:
            doc.add_paragraph(text, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph(
        "Michael Borck · michael.borck@curtin.edu.au · everything also at slinkr.link/udl")
    foot.alignment = WD_ALIGN_PARAGRAPH.LEFT
    foot.runs[0].font.size = Pt(9)

    doc.save(path)


def main() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir()

    for bundle in BUNDLES:
        folder = OUT / bundle["dir"]
        folder.mkdir()

        cover = folder / "READ_ME_FIRST.docx"
        write_cover(bundle, cover)

        missing = [f for f in bundle["files"] if not (ROOT / f).exists()]
        if missing:
            raise SystemExit(f"missing from {bundle['dir']}: {', '.join(missing)}")
        for name in bundle["files"]:
            # copy, not copy2: preserving xattrs on this exFAT volume spawns
            # AppleDouble ._ sidecars, which then land in the zip as junk.
            shutil.copy(ROOT / name, folder / name)

        archive = OUT / f"{bundle['dir']}.zip"
        with zipfile.ZipFile(archive, "w", zipfile.ZIP_DEFLATED) as zf:
            # Flat inside the zip - one folder, no nesting to click through.
            for item in sorted(folder.iterdir()):
                if item.name.startswith("._") or item.name == ".DS_Store":
                    continue
                zf.write(item, f"{bundle['dir']}/{item.name}")
        kb = archive.stat().st_size // 1024
        print(f"  {archive.name}  ({len(bundle['files']) + 1} files, {kb} kB)")

    # Loose, deliberately: this is the one he is asked to mark up.
    shutil.copy2(ROOT / "UDL3_Coverage_Map_for_review.docx", OUT)
    kb = (OUT / "UDL3_Coverage_Map_for_review.docx").stat().st_size // 1024
    print(f"  UDL3_Coverage_Map_for_review.docx  (loose, {kb} kB)")


if __name__ == "__main__":
    main()
